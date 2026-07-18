from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "cv"
OUTPUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUTPUT / "Tran_Thuan_Hoan_Research_CV.docx"

NAVY = "0F172A"
BLUE = "1D4ED8"
SLATE = "475569"
LIGHT = "CBD5E1"
TEAL = "0F766E"


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn(f"w:{key}"), str(value))


def set_paragraph_bottom_border(paragraph, color=LIGHT, size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=False):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    properties.append(color_element)
    if underline:
        underline_element = OxmlElement("w:u")
        underline_element.set(qn("w:val"), "single")
        properties.append(underline_element)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_run(run, size=None, bold=None, color=None, italic=None, font="Aptos"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic
    return run


def add_section_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.style = "CV Section"
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_bottom_border(paragraph, BLUE, "8")
    set_run(paragraph.add_run(text.upper()), 10.2, True, NAVY)
    return paragraph


def add_entry_header(document, title, right_text=None, subtitle=None, url=None):
    paragraph = document.add_paragraph()
    paragraph.style = "CV Entry"
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(7.05), WD_TAB_ALIGNMENT.RIGHT
    )
    set_run(paragraph.add_run(title), 9.8, True, NAVY)
    if url:
        paragraph.add_run("  ")
        add_hyperlink(paragraph, "Repository ↗", url, TEAL)
    if right_text:
        paragraph.add_run("\t")
        set_run(paragraph.add_run(right_text), 8.8, True, SLATE)
    if subtitle:
        sub = document.add_paragraph()
        sub.style = "CV Compact"
        sub.paragraph_format.keep_with_next = True
        set_run(sub.add_run(subtitle), 8.9, False, SLATE, True)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="CV Bullet")
        paragraph.paragraph_format.keep_together = True
        set_run(paragraph.add_run("• "), 9.35, True, BLUE)
        set_run(paragraph.add_run(item), 9.35, False, NAVY)


def add_label_line(document, label, text):
    paragraph = document.add_paragraph(style="CV Compact")
    set_run(paragraph.add_run(f"{label}: "), 9.1, True, NAVY)
    set_run(paragraph.add_run(text), 9.1, False, NAVY)
    return paragraph


document = Document()
section = document.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(0.48)
section.bottom_margin = Inches(0.48)
section.left_margin = Inches(0.58)
section.right_margin = Inches(0.58)
section.header_distance = Inches(0.2)
section.footer_distance = Inches(0.22)

styles = document.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
normal.font.size = Pt(9.35)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(2.6)
normal.paragraph_format.line_spacing = 1.06

section_style = styles.add_style("CV Section", WD_STYLE_TYPE.PARAGRAPH)
section_style.paragraph_format.space_before = Pt(8)
section_style.paragraph_format.space_after = Pt(4)

entry_style = styles.add_style("CV Entry", WD_STYLE_TYPE.PARAGRAPH)
entry_style.paragraph_format.space_before = Pt(4)
entry_style.paragraph_format.space_after = Pt(0.5)

compact_style = styles.add_style("CV Compact", WD_STYLE_TYPE.PARAGRAPH)
compact_style.paragraph_format.space_after = Pt(2)
compact_style.paragraph_format.line_spacing = 1.03

bullet_style = styles.add_style("CV Bullet", WD_STYLE_TYPE.PARAGRAPH)
bullet_style.paragraph_format.left_indent = Inches(0.17)
bullet_style.paragraph_format.first_line_indent = Inches(-0.13)
bullet_style.paragraph_format.space_after = Pt(2)
bullet_style.paragraph_format.line_spacing = 1.04
bullet_style._element.get_or_add_pPr().append(OxmlElement("w:keepNext"))

# Header
name = document.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(1)
set_run(name.add_run("TRAN THUAN HOAN"), 23, True, NAVY, font="Aptos Display")

identity = document.add_paragraph()
identity.alignment = WD_ALIGN_PARAGRAPH.CENTER
identity.paragraph_format.space_after = Pt(3)
set_run(
    identity.add_run(
        "UNDERGRADUATE RESEARCHER  ·  COMPUTATIONAL ECONOMICS  ·  QUANTITATIVE METHODS"
    ),
    8.6,
    True,
    BLUE,
)

contact = document.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(7)
set_run(contact.add_run("Hanoi, Vietnam  ·  "), 8.7, False, SLATE)
add_hyperlink(contact, "tthoan309@gmail.com", "mailto:tthoan309@gmail.com")
set_run(contact.add_run("  ·  "), 8.7, False, SLATE)
add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com/in/1marcuz1/")
set_run(contact.add_run("  ·  "), 8.7, False, SLATE)
add_hyperlink(contact, "GitHub", "https://github.com/tthoan309-ux")
set_run(contact.add_run("  ·  "), 8.7, False, SLATE)
add_hyperlink(
    contact,
    "ORCID 0009-0007-6229-8585",
    "https://orcid.org/0009-0007-6229-8585",
)
set_paragraph_bottom_border(contact, NAVY, "12")

add_section_heading(document, "Research Profile")
profile = document.add_paragraph(style="CV Compact")
profile.paragraph_format.space_after = Pt(3)
set_run(
    profile.add_run(
        "International Economics undergraduate combining econometrics, machine learning, data engineering, and research software to study complex economic systems. Builds reproducible evidence pipelines—from data acquisition and normalization to modeling, interpretation, and evaluation—with current interests in energy transition, green growth, financial risk, trade networks, and AI for economics."
    ),
    9.45,
    False,
    NAVY,
)

add_section_heading(document, "Education")
add_entry_header(
    document,
    "Foreign Trade University (FTU)",
    "Expected 2027",
    "Bachelor of International Economics · High Quality Program · GPA: 3.72 / 4.00",
)
add_label_line(
    document,
    "Selected coursework",
    "Advanced Mathematics (10.0); Probability & Statistics (9.3); Econometrics I (9.5); Econometrics II (10.0)",
)
add_label_line(document, "English", "IELTS 7.0")

add_section_heading(document, "Current Working Paper")
add_entry_header(
    document,
    "Dynamic Relationship between Energy Transition and Green Growth",
    "In development",
    "The Moderating Role of Public Attention",
)
add_bullets(
    document,
    [
        "Developing a dynamic empirical framework to test how energy transition shapes green growth and when public attention changes the strength of that relationship.",
        "Current stage: literature review and data construction; model development and writing follow measurement validation.",
    ],
)

add_section_heading(document, "Selected Research Systems")
add_entry_header(
    document,
    "Financial Distress Prediction Pipeline",
    "2026 · In development",
    "End-to-end financial research infrastructure and explainable early-warning modeling",
    "https://github.com/tthoan309-ux/financialdistress",
)
add_bullets(
    document,
    [
        "Automates CafeF crawling and financial-statement collection; extracts irregular disclosures through PDF parsing and OCR with traceable source metadata.",
        "Normalizes accounting labels, units, periods, and identities into a versioned firm-level dataset; engineers liquidity, leverage, profitability, cash-flow, and trajectory features.",
        "Evaluates distress classifiers with temporal validation, rare-event metrics, calibration, leakage checks, and SHAP explanations; preserves extraction and modeling decisions as auditable research artifacts.",
    ],
)

add_entry_header(
    document,
    "ESG–RAG Index",
    "2026 · Prototype",
    "Evidence-grounded retrieval infrastructure for corporate ESG disclosures",
    "https://github.com/tthoan309-ux/ESG-RAG-index",
)
add_bullets(
    document,
    [
        "Builds a layout-aware ingestion, parsing, metadata, chunking, and retrieval pipeline for sustainability reports while preserving issuer, page, period, and extraction-version provenance.",
        "Prioritizes citation auditability and retrieval evaluation over unsourced generation, enabling reproducible comparison across firms and reporting periods.",
    ],
)

add_entry_header(
    document,
    "Macroeconomic Crisis Early Warning System",
    "2025 · Completed",
    "Explainable machine-learning system for policy monitoring",
)
add_bullets(
    document,
    [
        "Frames crisis detection as a rare-event classification problem and compares linear and nonlinear models using discrimination, calibration, and class-sensitive metrics.",
        "Uses SHAP to inspect system-level vulnerability drivers and country-level alerts, connecting predictive performance with policy-relevant interpretation.",
    ],
)

document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

add_section_heading(document, "Additional Research Evidence")
add_entry_header(
    document,
    "Economic Sanctions & Trade Network Reconfiguration",
    "2026 · Research design",
    "International trade, panel econometrics, and network science",
)
add_bullets(
    document,
    [
        "Combines product-level bilateral trade flows, sanctions events, gravity-style estimation, and network measures to separate lost trade from rerouting and partner substitution.",
    ],
)

add_entry_header(
    document,
    "Inventory Optimization under Demand Uncertainty",
    "2026 · Model developed",
    "Stochastic decision science and simulation-based policy evaluation",
    "https://github.com/tthoan309-ux/inventory_project",
)
add_bullets(
    document,
    [
        "Compares replenishment policies across Monte Carlo demand regimes and stress-tests service, holding-cost, shortage-risk, volatility, and lead-time assumptions.",
    ],
)

add_entry_header(
    document,
    "Vietnam–EU Monthly CN8 Trade Dataset",
    "2026 · Construction",
    "Reproducible product-level trade research infrastructure",
)
add_bullets(
    document,
    [
        "Automates acquisition, code concordance, schema validation, unit harmonization, duplicate checks, and versioned release generation for monthly CN8 trade flows.",
    ],
)

add_section_heading(document, "Research Experience & Leadership")
add_entry_header(
    document,
    "Vice President · International Economics Union (IEU)",
    "Current",
)
add_bullets(
    document,
    [
        "Coordinate research-oriented academic activities and interdisciplinary collaboration for the International Economics student community.",
    ],
)

add_entry_header(
    document,
    "Research & Development Department · Young Researchers Club, FTU",
    "Current",
)
add_bullets(
    document,
    [
        "Develop research concepts, review analytical approaches, and contribute to workshops, seminars, and student-led empirical research activities.",
    ],
)

add_entry_header(
    document,
    "Research Department · Global Economics Club, FTU",
    "Current",
)
add_bullets(
    document,
    [
        "Support evidence synthesis, quantitative analysis, and collaborative discussion of policy-relevant economic questions.",
    ],
)

add_entry_header(
    document,
    "Data Analysis Division · G'Contest Organizing Committee",
    "2025",
)
add_bullets(
    document,
    [
        "Supported a university-level data analysis competition with 500+ participants; coordinated analytical and operational tasks across the organizing team.",
    ],
)

add_section_heading(document, "Methods & Research Tools")
methods = document.add_table(rows=1, cols=2)
methods.autofit = False
methods.columns[0].width = Inches(3.55)
methods.columns[1].width = Inches(3.55)
methods.rows[0].cells[0].width = Inches(3.55)
methods.rows[0].cells[1].width = Inches(3.55)
methods.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
methods.rows[0].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
for cell in methods.rows[0].cells:
    cell.margin_top = 0
    cell.margin_bottom = 0
    set_cell_border(
        cell,
        top={"val": "nil"},
        left={"val": "nil"},
        bottom={"val": "nil"},
        right={"val": "nil"},
    )

left = methods.rows[0].cells[0]
left.text = ""
for label, text in [
    ("Econometrics", "Panel data; fixed/random effects; PPML; gravity models; specification diagnostics; quantitative and policy analysis"),
    ("Machine learning", "Logistic regression; random forest; XGBoost; regression; SHAP; calibration; temporal and comparative evaluation"),
    ("Programming", "Python; SQL; Stata; Pandas; NumPy; Scikit-learn"),
]:
    p = left.add_paragraph() if left.paragraphs[0].text else left.paragraphs[0]
    set_run(p.add_run(f"{label}: "), 8.8, True, NAVY)
    set_run(p.add_run(text), 8.8, False, NAVY)
    p.paragraph_format.space_after = Pt(3)

right = methods.rows[0].cells[1]
right.text = ""
for label, text in [
    ("Research infrastructure", "Web crawling; OCR; PDF parsing; data normalization; validation pipelines; versioned dataset construction"),
    ("Optimization & networks", "ALNS; Monte Carlo simulation; network centrality; community detection; graph embeddings"),
    ("Decision support", "Power BI; DAX; Matplotlib; KPI design; interactive dashboards; data storytelling"),
]:
    p = right.add_paragraph() if right.paragraphs[0].text else right.paragraphs[0]
    set_run(p.add_run(f"{label}: "), 8.8, True, NAVY)
    set_run(p.add_run(text), 8.8, False, NAVY)
    p.paragraph_format.space_after = Pt(3)

# Footer with page number field
footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(footer_p.add_run("TRAN THUAN HOAN  ·  RESEARCH CV  ·  "), 7.5, False, SLATE)
field_begin = OxmlElement("w:fldChar")
field_begin.set(qn("w:fldCharType"), "begin")
field_instruction = OxmlElement("w:instrText")
field_instruction.set(qn("xml:space"), "preserve")
field_instruction.text = " PAGE "
field_end = OxmlElement("w:fldChar")
field_end.set(qn("w:fldCharType"), "end")
footer_p.add_run()._r.append(field_begin)
footer_p.add_run()._r.append(field_instruction)
footer_p.add_run()._r.append(field_end)

document.core_properties.title = "Tran Thuan Hoan — Research CV"
document.core_properties.subject = "Computational Economics and Quantitative Research"
document.core_properties.author = "Tran Thuan Hoan"
document.core_properties.keywords = (
    "econometrics, machine learning, computational economics, data engineering, research"
)

document.save(DOCX_PATH)
print(DOCX_PATH)
